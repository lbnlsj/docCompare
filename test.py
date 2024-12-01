from docx import Document
import PyPDF2
import difflib
import logging
import traceback
import os
from typing import List, Dict, Tuple
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.opc.oxml import parse_xml

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def load_from_xml_v2(baseURI, rels_item_xml):
    """返回加载了关系的 _SerializedRelationships 实例"""
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ('../NULL', 'NULL'):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels


_SerializedRelationships.load_from_xml = load_from_xml_v2


def extract_pdf_text(pdf_path: str) -> List[str]:
    """从PDF文件中提取文本内容并返回行列表"""
    try:
        lines = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                page_lines = [line.strip() for line in text.split('\n') if line.strip()]
                lines.extend(page_lines)
        return lines
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def compare_docx_files(file1_path: str, file2_path: str) -> Dict:
    """比较两个DOCX文件并返回差异结果"""
    try:
        doc1 = Document(file1_path)
        doc2 = Document(file2_path)

        lines1 = []
        lines2 = []

        for para in doc1.paragraphs:
            if para.text.strip():
                lines1.append(para.text)

        for para in doc2.paragraphs:
            if para.text.strip():
                lines2.append(para.text)

        return compare_text_lines(lines1, lines2)

    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def compare_pdf_files(file1_path: str, file2_path: str) -> Dict:
    """比较两个PDF文件并返回差异结果"""
    try:
        lines1 = extract_pdf_text(file1_path)
        lines2 = extract_pdf_text(file2_path)
        return compare_text_lines(lines1, lines2)
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def is_similar_text(text1: str, text2: str) -> bool:
    """判断两段文本是否相似"""
    # 计算字符级别的相似度
    matcher = difflib.SequenceMatcher(None, text1, text2)
    similarity = matcher.ratio()
    # 如果相似度大于阈值，认为是相似文本
    return similarity > 0.6


def find_modifications(lines1: List[str], lines2: List[str]) -> List[Dict]:
    """查找并标记修改的内容，采用更严格的匹配规则"""
    changes = []
    line_number = 1

    # 使用 SequenceMatcher 进行整体序列比较
    sm = difflib.SequenceMatcher(None, lines1, lines2)

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            # 处理相同的行
            for line in lines1[i1:i2]:
                changes.append({
                    'type': 'unchanged',
                    'content': line,
                    'line_number': line_number
                })
                line_number += 1

        elif tag == 'replace' and (i2 - i1) == (j2 - j1):
            # 只有当替换的行数相同，且内容相似度高时才视为修改
            for old_line, new_line in zip(lines1[i1:i2], lines2[j1:j2]):
                # 计算行级别的相似度
                similarity = difflib.SequenceMatcher(None, old_line, new_line).ratio()
                if similarity > 0.8:  # 提高相似度阈值
                    changes.append({
                        'type': 'modification',
                        'content': {
                            'old': old_line,
                            'new': new_line
                        },
                        'line_number': line_number
                    })
                else:
                    # 如果相似度不够高，作为独立的删除和添加处理
                    changes.append({
                        'type': 'deletion',
                        'content': old_line,
                        'line_number': line_number
                    })
                    changes.append({
                        'type': 'addition',
                        'content': new_line,
                        'line_number': line_number
                    })
                line_number += 1
        else:
            # 处理删除的行
            for line in lines1[i1:i2]:
                changes.append({
                    'type': 'deletion',
                    'content': line,
                    'line_number': line_number
                })
                line_number += 1

            # 处理添加的行
            for line in lines2[j1:j2]:
                changes.append({
                    'type': 'addition',
                    'content': line,
                    'line_number': line_number
                })
                line_number += 1

    return changes


def compare_text_lines(lines1: List[str], lines2: List[str]) -> Dict:
    """比较两个文本行列表并返回差异结果"""
    try:
        # 预处理：移除空行和空白字符
        lines1 = [line.strip() for line in lines1 if line.strip()]
        lines2 = [line.strip() for line in lines2 if line.strip()]

        changes = find_modifications(lines1, lines2)

        # 统计更改
        stats = {
            'additions': len([c for c in changes if c['type'] == 'addition']),
            'deletions': len([c for c in changes if c['type'] == 'deletion']),
            'modifications': len([c for c in changes if c['type'] == 'modification']),
            'unchanged': len([c for c in changes if c['type'] == 'unchanged'])
        }

        return {
            'success': True,
            'changes': changes,
            'stats': stats,
            'total_lines': len(changes)
        }

    except Exception as e:
        logger.error(f"Error in compare_text_lines: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def compare_text_lines(lines1: List[str], lines2: List[str]) -> Dict:
    """比较两个文本行列表并返回差异结果"""
    try:
        changes = find_modifications(lines1, lines2)

        # 统计更改
        stats = {
            'additions': len([c for c in changes if c['type'] == 'addition']),
            'deletions': len([c for c in changes if c['type'] == 'deletion']),
            'modifications': len([c for c in changes if c['type'] == 'modification']),
            'unchanged': len([c for c in changes if c['type'] == 'unchanged'])
        }

        return {
            'success': True,
            'changes': changes,
            'stats': stats,
            'total_lines': len(changes)
        }

    except Exception as e:
        logger.error(f"Error in compare_text_lines: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def run_comparison_tests():
    """运行文档比较测试"""
    try:
        # 测试文件路径
        test_files = {
            'docx': ('data/t1.docx', 'data/t2.docx'),
            'pdf': ('data/t3.pdf', 'data/t4.pdf')
        }

        print("开始文档比较测试...\n")

        # DOCX文件测试
        print("=== DOCX文件比较测试 ===")
        docx_result = compare_docx_files(*test_files['docx'])
        if docx_result['success']:
            print("DOCX比较成功:")
            print(f"- 添加: {docx_result['stats']['additions']} 处")
            print(f"- 删除: {docx_result['stats']['deletions']} 处")
            print(f"- 修改: {docx_result['stats']['modifications']} 处")
            print(f"- 未改: {docx_result['stats']['unchanged']} 处")
            print("\n具体变更:")
            for change in docx_result['changes']:
                if change['type'] == 'modification':
                    print(
                        f"[修改] 行 {change['line_number']}: {change['content']['old']} -> {change['content']['new']}")
                elif change['type'] == 'addition':
                    print(f"[添加] 行 {change['line_number']}: {change['content']}")
                elif change['type'] == 'deletion':
                    print(f"[删除] 行 {change['line_number']}: {change['content']}")
        else:
            print(f"DOCX比较失败: {docx_result.get('error', '未知错误')}")

        print("\n=== PDF文件比较测试 ===")
        pdf_result = compare_pdf_files(*test_files['pdf'])
        if pdf_result['success']:
            print("PDF比较成功:")
            print(f"- 添加: {pdf_result['stats']['additions']} 处")
            print(f"- 删除: {pdf_result['stats']['deletions']} 处")
            print(f"- 修改: {pdf_result['stats']['modifications']} 处")
            print(f"- 未改: {pdf_result['stats']['unchanged']} 处")
            print("\n具体变更:")
            for change in pdf_result['changes']:
                if change['type'] == 'modification':
                    print(
                        f"[修改] 行 {change['line_number']}: {change['content']['old']} -> {change['content']['new']}")
                elif change['type'] == 'addition':
                    print(f"[添加] 行 {change['line_number']}: {change['content']}")
                elif change['type'] == 'deletion':
                    print(f"[删除] 行 {change['line_number']}: {change['content']}")
        else:
            print(f"PDF比较失败: {pdf_result.get('error', '未知错误')}")

    except Exception as e:
        print(f"测试过程出错: {str(e)}")
        print("详细错误信息:")
        traceback.print_exc()


if __name__ == '__main__':
    run_comparison_tests()