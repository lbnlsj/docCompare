from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import shutil
from docx import Document
import logging
import difflib
import traceback
import os
from typing import List, Dict, Tuple
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.opc.oxml import parse_xml


def load_from_xml_v2(baseURI, rels_item_xml):
    """返回加载了关系的 _SerializedRelationships 实例。"""
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ('../NULL', 'NULL'):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels


_SerializedRelationships.load_from_xml = load_from_xml_v2

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import shutil
from docx import Document
import logging
import difflib
import traceback
import os
import PyPDF2
from typing import List, Dict, Tuple
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.opc.oxml import parse_xml


def load_from_xml_v2(baseURI, rels_item_xml):
    """返回加载了关系的 _SerializedRelationships 实例。"""
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ('../NULL', 'NULL'):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels


_SerializedRelationships.load_from_xml = load_from_xml_v2


def extract_pdf_text(pdf_path: str) -> List[str]:
    """
    从PDF文件中提取文本内容并返回行列表
    """
    try:
        lines = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                # 按行分割文本
                page_lines = [line.strip() for line in text.split('\n') if line.strip()]
                lines.extend(page_lines)
        return lines
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def compare_docx_files(file1_path: str, file2_path: str) -> Dict:
    """
    比较两个DOCX文件并返回差异结果
    """
    try:
        # 提取文档内容
        doc1 = Document(file1_path)
        doc2 = Document(file2_path)

        # 将文档内容转换为行列表
        lines1 = []
        lines2 = []

        # 处理第一个文档
        for para in doc1.paragraphs:
            if para.text.strip():  # 忽略空行
                lines1.append(para.text)

        # 处理第二个文档
        for para in doc2.paragraphs:
            if para.text.strip():  # 忽略空行
                lines2.append(para.text)

        return compare_text_lines(lines1, lines2)

    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def compare_pdf_files(file1_path: str, file2_path: str) -> Dict:
    """
    比较两个PDF文件并返回差异结果
    """
    try:
        # 提取PDF文本内容
        lines1 = extract_pdf_text(file1_path)
        lines2 = extract_pdf_text(file2_path)

        return compare_text_lines(lines1, lines2)

    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def compare_text_lines(lines1: List[str], lines2: List[str]) -> Dict:
    """
    比较两个文本行列表并返回差异结果
    """
    try:
        # 使用difflib进行比较
        diff = list(difflib.ndiff(lines1, lines2))

        # 解析差异
        changes = []
        for i, line in enumerate(diff):
            if line.startswith('- '):
                # 删除的内容
                changes.append({
                    'type': 'deletion',
                    'content': line[2:],
                    'line_number': i + 1
                })
            elif line.startswith('+ '):
                # 添加的内容
                changes.append({
                    'type': 'addition',
                    'content': line[2:],
                    'line_number': i + 1
                })
            elif line.startswith('? '):
                # 指示具体的修改位置，可以忽略
                continue
            else:
                # 未修改的内容
                changes.append({
                    'type': 'unchanged',
                    'content': line[2:],
                    'line_number': i + 1
                })

        # 统计更改
        stats = {
            'additions': len([c for c in changes if c['type'] == 'addition']),
            'deletions': len([c for c in changes if c['type'] == 'deletion']),
            'unchanged': len([c for c in changes if c['type'] == 'unchanged'])
        }

        return {
            'success': True,
            'changes': changes,
            'stats': stats,
            'total_lines': len(changes)
        }

    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


@app.route('/compare', methods=['POST'])
def compare_files():
    """处理文件比较请求"""
    try:
        # 获取请求中的文件信息
        data = request.get_json()
        if not data or 'file1' not in data or 'file2' not in data:
            return jsonify({'error': 'Both files are required'}), 400

        file1_path = os.path.join('static', 'data', data['file1'])
        file2_path = os.path.join('static', 'data', data['file2'])

        # 检查文件是否存在
        if not os.path.exists(file1_path) or not os.path.exists(file2_path):
            return jsonify({'error': 'One or both files not found'}), 404

        # 获取文件扩展名
        file1_ext = os.path.splitext(file1_path)[1].lower()
        file2_ext = os.path.splitext(file2_path)[1].lower()

        # 检查文件类型是否匹配
        if file1_ext != file2_ext:
            return jsonify({'error': 'Files must be of the same type'}), 400

        # 根据文件类型选择比较方法
        if file1_ext == '.docx':
            result = compare_docx_files(file1_path, file2_path)
        elif file1_ext == '.pdf':
            result = compare_pdf_files(file1_path, file2_path)
        else:
            return jsonify({'error': 'Unsupported file type'}), 400

        if not result['success']:
            return jsonify({'error': result['error']}), 500

        # 返回比较结果
        return jsonify({
            'success': True,
            'comparison': result,
            'file1': data['file1'],
            'file2': data['file2']
        })

    except Exception as e:
        stack_trace = traceback.format_exc()
        error_msg = f"Error in compare_files: {str(e)}\nStack trace:\n{stack_trace}"
        logger.error(error_msg)
        return jsonify({'error': str(e)}), 500


def setup_static_files():
    """设置静态文件和目录结构"""
    try:
        # 创建必要的目录
        directories = ['static', 'static/data', 'data', 'templates']
        for directory in directories:
            os.makedirs(directory, exist_ok=True)
            logger.info(f"Created or verified directory: {directory}")

        # 复制默认文件到静态目录
        default_files = ['t1.docx', 't2.docx']
        for file in default_files:
            source = os.path.join('data', file)
            destination = os.path.join('static', 'data', file)
            if os.path.exists(source):
                shutil.copy2(source, destination)
                logger.info(f"Copied {file} to static/data directory")
            else:
                logger.warning(f"Source file not found: {source}")
    except Exception as e:
        logger.error(f"Error in setup_static_files: {str(e)}")
        raise


@app.route('/')
def index():
    """渲染主页"""
    return render_template('index.html')


@app.route('/static/<path:path>')
def send_static(path):
    """处理静态文件请求"""
    return send_from_directory('static', path)


@app.route('/static/data/<filename>')
def serve_default_file(filename):
    """服务默认文件"""
    try:
        return send_from_directory(os.path.join('static', 'data'), filename)
    except Exception as e:
        logger.error(f"Error serving file {filename}: {str(e)}")
        return jsonify({'error': 'File not found'}), 404


@app.route('/upload', methods=['POST'])
def upload_file():
    """处理文件上传"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        # 获取文件扩展名
        ext = os.path.splitext(file.filename)[1].lower()

        # 检查文件类型
        if ext not in ['.pdf', '.docx']:
            return jsonify({'error': 'Invalid file type'}), 400

        # 保存文件到临时目录
        temp_path = os.path.join('static', 'temp', file.filename)
        os.makedirs(os.path.dirname(temp_path), exist_ok=True)
        file.save(temp_path)

        return jsonify({
            'success': True,
            'filename': file.filename,
            'path': f'/static/temp/{file.filename}'
        })

    except Exception as e:
        logger.error(f"Error in upload_file: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/cleanup', methods=['POST'])
def cleanup_temp_files():
    """清理临时文件"""
    try:
        temp_dir = os.path.join('static', 'temp')
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            os.makedirs(temp_dir)
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error in cleanup_temp_files: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.errorhandler(404)
def not_found_error(error):
    """处理404错误"""
    return jsonify({'error': 'Not found'}), 404


@app.errorhandler(500)
def internal_error(error):
    """处理500错误"""
    logger.error(f"Internal server error: {str(error)}")
    return jsonify({'error': 'Internal server error'}), 500


# 初始化设置
setup_static_files()

if __name__ == '__main__':
    try:
        # 确保临时目录存在
        os.makedirs(os.path.join('static', 'temp'), exist_ok=True)

        logger.info("Starting Flask application...")
        app.run(debug=True, host='0.0.0.0', port=5000)
    except Exception as e:
        logger.error(f"Failed to start application: {str(e)}")
